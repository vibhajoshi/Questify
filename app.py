from flask import Flask, request, jsonify, send_file
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from docx import Document
from pdf2image import convert_from_path
import pytesseract
import os
from werkzeug.utils import secure_filename
from flask import render_template
import requests  # For making HTTP requests to Ollama's local API
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text

# Initialize Flask app
app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20 MB

# Create a temporary folder for storing files
UPLOAD_FOLDER = "tmp"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Ollama API configuration
OLLAMA_API_URL = "http://localhost:11434/api/generate"  # Default Ollama API endpoint
OLLAMA_MODEL = "llama3.2"  # Replace with your installed model (e.g., "llama3.2")



def extract_text_from_docx(file, start_page=None, end_page=None):
    """Extract text from a DOCX file, optionally within a page range."""
    try:
        doc = Document(file)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Simulate page range by splitting paragraphs into chunks
        if start_page is not None and end_page is not None:
            chunk_size = len(paragraphs) // 10  # Assume 10 "pages" for simplicity
            start_index = (start_page - 1) * chunk_size
            end_index = end_page * chunk_size
            paragraphs = paragraphs[start_index:end_index]
        
        # Add simulated page numbers
        text = ""
        for i, para in enumerate(paragraphs):
            page_num = (i // chunk_size) + 1 if chunk_size > 0 else 1
            text += f"Page {page_num}:\n{para}\n\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


from bs4 import BeautifulSoup

def extract_text_from_html(file, start_page=None, end_page=None):
    """Extract text from an HTML file, optionally within a page range."""
    try:
        file.seek(0)
        content = file.read().decode('utf-8')
        soup = BeautifulSoup(content, "html.parser")
        text = soup.get_text(separator="\n")
        
        # Simulate page range by splitting text into chunks
        if start_page is not None and end_page is not None:
            lines = text.splitlines()
            chunk_size = len(lines) // 10  # Assume 10 "pages" for simplicity
            start_index = (start_page - 1) * chunk_size
            end_index = end_page * chunk_size
            lines = lines[start_index:end_index]
        
        # Add simulated page numbers
        text = ""
        for i, line in enumerate(lines):
            page_num = (i // chunk_size) + 1 if chunk_size > 0 else 1
            text += f"Page {page_num}:\n{line}\n\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from HTML: {e}")
        return None

"""
def extract_text_from_html(file):
    #Extract text from an HTML file.
    try:
        content = file.read().decode("utf-8")  # Decode the file content
        soup = BeautifulSoup(content, "html.parser")
        return soup.get_text(separator="\n").strip()  # Use separator to preserve line breaks
    except Exception as e:
        print(f"Error extracting text from HTML: {e}")
        return None
"""

def extract_text_from_pdf(file_path, start_page=None, end_page=None):
    """Extract text from PDF, optionally within a page range."""
    try:
        from pdfminer.high_level import extract_text

        if start_page is not None and end_page is not None:
            text = ""
            for page_num in range(start_page - 1, end_page):  # pdfminer uses 0-based indexing
                text += f"Page {page_num + 1}:\n"  # Add page number
                text += extract_text(file_path, page_numbers=[page_num]) + "\n\n"
        else:
            text = extract_text(file_path)
        
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


def generate_questions_with_ollama(text, num_questions, difficulty, marks, start_page=None, end_page=None):
    """Generate questions using Ollama's local API based on difficulty and marks."""
    try:
        # Define the prompt based on difficulty level and marks
        if difficulty == "basic":
            if marks == "1":
                prompt = f"""
                Generate exactly {num_questions} basic-level 1 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Generate multiple-choice questions (MCQs) or fill-in-the-blank questions.
                3. For MCQs, provide the question followed by 4 options (A, B, C, D) and the correct answer.
                4. For fill-in-the-blank questions, provide the question with a blank (_____) and the correct answer.
                5. Ensure each question is self-contained and does not reference other questions.
                6. Format the output as follows:
                    - For MCQs:
                        Q1. [Question]
                        A) [Option A]
                        B) [Option B]
                        C) [Option C]
                        D) [Option D]
                        Answer: [Correct Option]
                    - For fill-in-the-blank:
                        Q1. [Question with a blank]
                        Answer: [Correct Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                """
            elif marks == "2":
                prompt = f"""
                Generate exactly {num_questions} basic-level 2 Mark questions based on the following text:
                
                {text}

                Rules:
                1. 1. Ensure each question is self-contained and does not reference other questions.
                2. Generate questions that require a 2-line answer.
                3. Questions should be of the type: "What is...", "How is...", "What are...", etc.
                4. Do not include MCQs or fill-in-the-blank questions.
                5. Ensure each question is self-contained and does not reference other questions.
                6. Format the output as follows:
                    Q1. [Question]
                    Answer: [2-line Answer]
                    Q2. [Question]
                    Answer: [2-line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [2-line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                """
            elif marks == "5":
                prompt = f"""
                Generate exactly {num_questions} basic-level 5 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Generate questions that require a 5-line answer.
                3. Questions should be of the type: "What is...", "Explain...", "Describe...", etc.
                4. Do not include MCQs or fill-in-the-blank questions.
                5. Ensure each question is self-contained and does not reference other questions.
                6. Format the output as follows:
                    Q1. [Question]
                    Answer: [5-line Answer]
                    Q2. [Question]
                    Answer: [5-line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [5-line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "10":
                prompt = f"""
                Generate exactly {num_questions} basic-level 10 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Generate questions that require a 10-line answer.
                3. Questions should be of the type: "What is...", "Explain in detail...", "Discuss...", etc.
                4. Do not include MCQs or fill-in-the-blank questions.
                5. Ensure each question is self-contained and does not reference other questions.
                6. Format the output as follows:
                    Q1. [Question]
                    Answer: [10-line Answer]
                    Q2. [Question]
                    Answer: [10-line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [10-line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                
                """
            else:
                return jsonify({"error": "Invalid marks value. Choose from '1', '2', '5', or '10'."}), 400
        elif difficulty == "intermediate":
            if marks == "1":
                prompt = f"""
                Generate exactly {num_questions} intermediate-level 1 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Remembering and Understanding.
                3. Generate descriptive questions that require short answers (2-3 lines).
                4. Questions should test basic knowledge and comprehension.
                5. Format the output as follows:
                    Q1. [Question]
                    Answer: [2-3 line Answer]
                    Q2. [Question]
                    Answer: [2-3 line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [2-3 line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "2":
                prompt = f"""
                Generate exactly {num_questions} intermediate-level 2 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Remembering and Understanding.
                3. Generate descriptive questions that require short answers (2-3 lines).
                4. Questions should test basic knowledge and comprehension.
                5. Format the output as follows:
                    Q1. [Question]
                    Answer: [2-3 line Answer]
                    Q2. [Question]
                    Answer: [2-3 line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [2-3 line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "5":
                prompt = f"""
                Generate exactly {num_questions} intermediate-level 5 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Applying and Analyzing.
                3. Generate descriptive questions that require detailed explanations (8-9 lines).
                4. Questions should test the ability to apply and analyze concepts.
                5. Format the output as follows:
                    Q1. [Question]
                    Answer: [8-9 line Answer]
                    Q2. [Question]
                    Answer: [8-9 line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [8-9 line Answer]
                6. Do not include any additional text, explanations, or notes.
                7. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "10":
                prompt = f"""
                Generate exactly {num_questions} intermediate-level 10 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Evaluating and Creating.
                3. Generate descriptive, discussion-based, or critical-thinking questions that require comprehensive answers (15-20 lines).
                4. Questions should test the ability to evaluate and create solutions or arguments.
                5. Format the output as follows:
                    Q1. [Question]
                    Answer: [15-20 line Answer]
                    Q2. [Question]
                    Answer: [15-20 line Answer]
                    ...
                    Q{num_questions}. [Question]
                    Answer: [15-20 line Answer]
                6. Do not include any additional text, explanations, or notes.
                7. Ensure the numbering starts from Q1 and is continuous.
                
                """
            else:
                return jsonify({"error": "Invalid marks value. Choose from '1', '2', '5', or '10'."}), 400

        elif difficulty == "complex":
            if marks == "1":
                prompt = f"""
                Generate exactly {num_questions} complex-level 1 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Applying and Analyzing.
                3. Generate complex questions that require concise answers (1-2 lines).
                4. Questions should test the ability to apply and analyze concepts.
                5. Format the output as follows:
                    Q1. [Complex Question]
                    Answer: [1-2 line Answer]
                    Q2. [Complex Question]
                    Answer: [1-2 line Answer]
                    ...
                    Q{num_questions}. [Complex Question]
                    Answer: [1-2 line Answer]
                6. Do not include any additional text, explanations, or notes.
                7. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "2":
                prompt = f"""
                Generate exactly {num_questions} complex-level 2 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Applying and Analyzing.
                3. Generate applied questions that require short explanations (4-5 lines).
                4. Questions should test the ability to apply and analyze concepts.
                5. Format the output as follows:
                    Q1. [Applied Question]
                    Answer: [4-5 line Answer]
                    Q2. [Applied Question]
                    Answer: [4-5 line Answer]
                    ...
                    Q{num_questions}. [Applied Question]
                    Answer: [4-5 line Answer]
                6. Do not include any additional text, explanations, or notes.
                7. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "5":
                prompt = f"""
                Generate exactly {num_questions} complex-level 5 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Evaluating and Creating.
                3. Generate out-of-the-box applied questions, including diagram-based questions (flowcharts, architectures, etc.).
                4. Questions should challenge students to think creatively and apply concepts in innovative ways.
                5. Provide detailed answers (10-15 lines) that explain the reasoning or solution.
                6. Format the output as follows:
                    Q1. [Applied Question]
                    Answer: [10-15 line Answer]
                    Q2. [Diagram-Based Question]
                    Answer: [10-15 line Answer]
                    ...
                    Q{num_questions}. [Applied Question]
                    Answer: [10-15 line Answer]
                7. Do not include any additional text, explanations, or notes.
                8. Ensure the numbering starts from Q1 and is continuous.
                
                """
            elif marks == "10":
                prompt = f"""
                Generate exactly {num_questions} complex-level 10 Mark questions based on the following text:
                
                {text}

                Rules:
                1. Ensure each question is self-contained and does not reference other questions.
                2. Focus on Bloom's Taxonomy levels: Evaluating and Creating.
                3. Generate complex, conceptualized questions with diagrams/flowcharts/architectures.
                4. Questions should require comprehensive answers (20-25 lines) that demonstrate deep understanding and critical thinking.
                5. Format the output as follows:
                    Q1. [Complex Conceptualized Question]
                    Answer: [20-25 line Answer]
                    Q2. [Diagram-Based Question]
                    Answer: [20-25 line Answer]
                    ...
                    Q{num_questions}. [Complex Conceptualized Question]
                    Answer: [20-25 line Answer]
                6. Do not include any additional text, explanations, or notes.
                7. Ensure the numbering starts from Q1 and is continuous.
                
                """
            else:
                return jsonify({"error": "Invalid marks value. Choose from '1', '2', '5', or '10'."}), 400
        else:
            return jsonify({"error": "Invalid difficulty level. Choose from 'basic', 'intermediate', or 'complex'."}), 400

        if start_page is not None and end_page is not None:
            prompt += f"\n\nNote: The questions should be generated from pages {start_page} to {end_page}."

        # Prepare the payload for Ollama API
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False  # Set to False to get a single response
        }

        # Call Ollama API and retry until the exact number of questions is generated
        max_retries = 8
        for attempt in range(max_retries):
            response = requests.post(OLLAMA_API_URL, json=payload)
            response.raise_for_status()  # Raise an error for bad status codes

            # Extract questions from the response
            if response.status_code == 200:
                response_data = response.json()
                response_text = response_data.get("response", "").strip()

                # Split the response into lines and filter out unnecessary lines
                lines = response_text.split("\n")
                filtered_lines = [
                    line.strip() for line in lines 
                    if line.strip() and not line.strip().lower().startswith(("here are", "questions based on", "rules:"))
                ]

                # Group questions and ensure continuous numbering
                questions = []
                current_question = ""
                question_number = 1  # Start numbering from 1

                for line in filtered_lines:
                    if line.startswith("Q"):
                        if current_question:
                            questions.append(current_question)
                        # Update the question number to ensure continuity
                        current_question = line
                        question_number += 1
                    else:
                        current_question += "\n" + line

                if current_question:
                    questions.append(current_question)

                 # Check if the number of questions matches the requested number
                if len(questions) == num_questions:
                    return questions  # Return only the requested number of questions
                elif len(questions) > num_questions:
                    return questions[:num_questions]  # Trim excess questions
            
            print(f"Attempt {attempt + 1}: Did not generate enough questions. Retrying...")

        print(f"Failed to generate {num_questions} questions after {max_retries} attempts.")
        return None

    except Exception as e:
        print(f"Error generating questions with Ollama: {e}")
        return None

def create_word_document(questions):
    """Create and save a Word document with generated questions."""
    try:
        doc = Document()
        doc.add_heading("Generated Questions", level=1)
        for question in questions:
            doc.add_paragraph(question)
        doc_path = os.path.join(UPLOAD_FOLDER, "generated_questions.docx")
        doc.save(doc_path)
        return doc_path
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None

@app.route("/generate-questions", methods=["POST"])
def generate_questions_api():
    """API endpoint to generate questions."""
    try:
        file = request.files.get("file")
        num_questions = int(request.form.get("num_questions", 5))
        difficulty = request.form.get("difficulty", "basic")
        marks = request.form.get("marks", "1")  # Default to 1 Mark if not provided
        start_page = request.form.get("start_page", type=int)
        end_page = request.form.get("end_page", type=int)

        if not file or file.filename == "":
            return jsonify({"error": "No file uploaded or file is empty."}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)  # Save the file temporarily

        # Check file format and extract text
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path, start_page, end_page)
        elif filename.endswith(".html"):
            text = extract_text_from_html(file, start_page, end_page)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(file, start_page, end_page)
        else:
            return jsonify({"error": "Unsupported file format. Only PDF, HTML, and DOCX are supported."}), 400
        if not text:
            return jsonify({"error": "Failed to extract text from the file."}), 400

        # Generate questions
        questions = generate_questions_with_ollama(text, num_questions, difficulty, marks, start_page, end_page)
        if not questions:
            return jsonify({"error": "Failed to generate questions. Please try again."}), 500

        # Create Word document
        doc_path = create_word_document(questions)
        if not doc_path:
            return jsonify({"error": "Failed to create Word document."}), 500

        # Clean up the temporary file
        if os.path.exists(file_path):
            os.remove(file_path)

        return jsonify({
            "questions": questions,
            "message": "Questions generated successfully!",
            "download_link": f"/download/{os.path.basename(doc_path)}"
        })
    except Exception as e:
        print(f"Unexpected error in /generate-questions: {e}")
        return jsonify({"error": "An unexpected error occurred. Please try again."}), 500

@app.route("/")
def home():
    """Render the home page with the upload form."""
    return render_template("index.html")

@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    """Download the generated Word document."""
    try:
        file_path = os.path.join(UPLOAD_FOLDER, secure_filename(filename))
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found."}), 404
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"Error downloading file: {e}")
        return jsonify({"error": "An error occurred while downloading the file."}), 500

if __name__ == "__main__":
    app.run(debug=True,port=5001)  # Enable debug mode